# -*- coding: utf-8 -*-
"""modules/*.json 하나에서 학생 배포용 HTML · 학생 워크북 · 교사 가이드 · 허브를 모두 생성한다."""
import json, io, os, glob, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # 이 스크립트의 상위 폴더
MODDIR = os.path.join(ROOT, "modules")
DOCDIR = os.path.join(ROOT, "워드")
TPL = os.path.join(ROOT, "_템플릿", "스텝러너_템플릿.html")
FONT = "맑은 고딕"
os.makedirs(DOCDIR, exist_ok=True)

# ---------------- docx 헬퍼 ----------------
def new_doc():
    doc = Document(); s = doc.sections[0]
    s.page_width, s.page_height = Cm(21), Cm(29.7)
    s.top_margin = s.bottom_margin = Cm(1.9); s.left_margin = s.right_margin = Cm(2.0)
    st = doc.styles["Normal"]; st.font.name = FONT; st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.space_after = Pt(4); st.paragraph_format.line_spacing = 1.25
    return doc

def shade(cell, hexcolor):
    el = OxmlElement("w:shd"); el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto"); el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)

def borders(cell, color="D7E2F5", sz=8):
    b = OxmlElement("w:tcBorders")
    for s in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:" + s)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color); b.append(e)
    cell._tc.get_or_add_tcPr().append(b)

def cellpad(cell, pt=6):
    m = OxmlElement("w:tcMar")
    for s in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:" + s); e.set(qn("w:w"), str(int(pt * 20)))
        e.set(qn("w:type"), "dxa"); m.append(e)
    cell._tc.get_or_add_tcPr().append(m)

def para(doc, text="", size=10.5, bold=False, color=None, sb=0, sa=4, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if color: r.font.color.rgb = RGBColor.from_string(color)
    return p

def hr(doc, color="C8D2E0"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(8)
    b = OxmlElement("w:pBdr"); e = OxmlElement("w:bottom")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "6"); e.set(qn("w:space"), "1")
    e.set(qn("w:color"), color); b.append(e); p._p.get_or_add_pPr().append(b)

def box(doc, lines, fill="F8FBFF", border="D7E2F5", size=10.5, bold_first=False):
    t = doc.add_table(rows=1, cols=1); t.autofit = False
    t.columns[0].width = Cm(17.0)
    c = t.cell(0, 0); c.width = Cm(17.0)
    shade(c, fill); borders(c, border); cellpad(c, 7); c.text = ""
    for i, ln in enumerate(lines):
        p = c.paragraphs[0] if i == 0 else c.add_paragraph()
        p.paragraph_format.space_after = Pt(1); p.paragraph_format.line_spacing = 1.3
        r = p.add_run(ln); r.font.size = Pt(size); r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        if bold_first and i == 0: r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def writebox(doc, h_cm=3.0, label=None, guide=""):
    if label: para(doc, label, size=9.5, bold=True, color="5B6675", sb=4, sa=2)
    if guide and "\n" in guide:
        for ln in guide.split("\n"):
            para(doc, ln if ln.strip() else " ", size=9, color="9AA4B2", sa=0, indent=0.3)
        h_cm = max(1.2, h_cm - 0.35 * len(guide.split("\n")))
    t = doc.add_table(rows=1, cols=1); t.autofit = False
    t.columns[0].width = Cm(17.0)
    row = t.rows[0]; row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST; row.height = Cm(h_cm)
    c = t.cell(0, 0); c.width = Cm(17.0); borders(c, "B9C4D6", 6); cellpad(c, 6)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def table(doc, rows, widths, header=True, size=10, rowheight=None):
    t = doc.add_table(rows=len(rows), cols=len(widths)); t.autofit = False
    for ci, w in enumerate(widths): t.columns[ci].width = Cm(w)
    for ri, row in enumerate(rows):
        if rowheight and ri > 0:
            t.rows[ri].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            t.rows[ri].height = Cm(rowheight)
        for ci, val in enumerate(row):
            c = t.cell(ri, ci); c.width = Cm(widths[ci])
            borders(c, "C8D2E0", 6); cellpad(c, 5)
            if header and ri == 0: shade(c, "EEF3FF")
            c.text = ""
            for li, ln in enumerate(str(val).split("\n")):
                p = c.paragraphs[0] if li == 0 else c.add_paragraph()
                p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.2
                r = p.add_run(ln); r.font.size = Pt(size); r.font.name = FONT
                r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
                if header and ri == 0: r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def head(doc, kicker, title, sub):
    para(doc, kicker, size=9.5, bold=True, color="2F6DF0", sa=1)
    para(doc, title, size=19, bold=True, sa=2)
    para(doc, sub, size=10.5, color="5B6675", sa=8)
    hr(doc)

def h2(doc, text):
    para(doc, text, size=14, bold=True, sb=10, sa=3); hr(doc, "DCE3EE")

def lbl(m, text):
    """인쇄물에서는 {{키}}를 [보이는 이름]으로 바꿔 학생이 무엇을 채울지 알게 한다."""
    if not text: return text
    names = {x["k"]: x["label"] for x in m.get("mats", [])}
    return re.sub(r"\{\{(\w+)\}\}", lambda mo: "[" + names.get(mo.group(1), mo.group(1)) + "]", text)

# ---------------- 학생용 워크북 ----------------
def build_workbook(m, path):
    doc = new_doc(); it = m.get("intro", {})
    head(doc, m["stage"] + " · 공통 과제 모듈 " + m["code"], m["title"],
         "AI와 한 스텝씩 같이 만드는 학습 워크북  |  " + m.get("subject", ""))
    table(doc, [["학반", "", "번호", "", "이름", ""]], [2.0, 3.6, 2.0, 2.4, 2.0, 5.0], header=False, size=10.5)

    para(doc, "이 모듈에서 하는 일", size=13, bold=True, sb=8, sa=3)
    box(doc, [it.get("lead", "")], fill="F7F9FC", border="C9D6EC")
    if it.get("phases"):
        table(doc, [["차시", "하는 일"]] + [[p.split("·")[0].strip(), "·".join(p.split("·")[1:]).strip()] for p in it["phases"]],
              [2.6, 14.4], size=10)
    if it.get("outcomeItems"):
        para(doc, "오늘 끝나면 이런 것이 나옵니다", size=13, bold=True, sb=8, sa=3)
        box(doc, [it.get("outcomeTitle", "")] + ["%d. %s" % (i + 1, x) for i, x in enumerate(it["outcomeItems"])],
            fill="FBFCFF", border="C9D6EC", bold_first=True)
    if it.get("rules"):
        para(doc, "시작 전 약속", size=13, bold=True, sb=8, sa=3)
        box(doc, ["%s %s" % ("①②③④⑤"[i], r) for i, r in enumerate(it["rules"])],
            fill="FFF8E6", border="F5DFA0")
    doc.add_page_break()

    cur = 0
    for st in m["steps"]:
        pn = m.get("phaseNames", {}).get(str(st.get("phase")), "")
        if st.get("phase") != cur:
            cur = st.get("phase")
            para(doc, "◆ " + pn, size=12, bold=True, color="2F6DF0", sb=2, sa=6)
        para(doc, st["no"] + "  " + st["title"], size=15, bold=True, sb=8, sa=2)
        hr(doc, "DCE3EE")
        if st.get("why"): para(doc, lbl(m, st["why"]), size=9.5, color="5B6675", sa=5)
        if st.get("noai"):
            box(doc, ["※ 이 스텝에서는 AI를 쓰지 않습니다. 내 머리와 내 손으로 합니다."],
                fill="FDF3F2", border="F3D3CE", size=10)
        if st.get("panel") and st["panel"].get("lines"):
            box(doc, ([lbl(m, st["panel"].get("title", ""))] if st["panel"].get("title") else [])
                     + [lbl(m, x) for x in st["panel"]["lines"]],
                fill="FBFCFF", border="C9D6EC", size=11, bold_first=bool(st["panel"].get("title")))
        if st.get("mats"):
            table(doc, [["항목", "내가 적을 내용"]] + [[x["label"], ""] for x in m["mats"]],
                  [5.0, 12.0], size=10, rowheight=0.8)
        if st.get("prompts"):
            all_nocopy = all(p.get("noCopy") for p in st["prompts"])
            head_nocopy = all_nocopy or bool(st["prompts"][0].get("noCopy"))
            t = "▶ 이렇게 하세요" if head_nocopy else "▶ AI에게 이렇게 보내세요"
            if len(st["prompts"]) > 1 and not all_nocopy:
                t += "   (둘 중 하나만 고르세요)" if st.get("pickOne") else "   (반드시 한 번에 하나씩)"
            para(doc, t, size=10.5, bold=True, color="1F53C4", sb=4, sa=3)
            for p in st["prompts"]:
                box(doc, [p.get("label", "")] + lbl(m, p["text"]).split("\n"),
                    fill="F8FBFF", border="D7E2F5", size=10, bold_first=True)
        if st.get("expect"):
            para(doc, "· 이런 답이 오면 정상입니다", size=9.5, bold=True, color="5B6675", sa=2)
            box(doc, lbl(m, st["expect"]).split("\n"), fill="FAFBFD", border="E2E6EC", size=9)
        if st.get("troubles"):
            para(doc, "· 이렇게 나왔다면? — 막혔을 때 보낼 문장", size=9.5, bold=True, color="B3590A", sa=2)
            for t in st["troubles"]:
                box(doc, ["[막혔을 때] " + t["when"]] + ["→ " + l for l in lbl(m, t["fix"]).split("\n")],
                    fill="FFFAF3", border="F2DFC4", size=9.5, bold_first=True)
        if st.get("record"):
            h = 7.5 if st["record"].get("big") else 3.6
            writebox(doc, h, "▷ 기록 — " + lbl(m, st["record"]["label"]),
                     guide=lbl(m, st["record"].get("ph", "")))
        if st.get("checks"):
            para(doc, "✔ 넘어가기 전 확인", size=9.5, bold=True, color="5B6675", sa=2)
            for c in st["checks"]: para(doc, "☐  " + lbl(m, c), size=10, sa=1, indent=0.3)
        doc.add_page_break()

    para(doc, "AI 사용 기록표", size=15, bold=True, sa=2); hr(doc, "DCE3EE")
    para(doc, "AI를 썼다는 사실은 숨기는 것이 아니라 밝히는 것입니다. 평가의 근거가 됩니다.",
         size=9.5, color="5B6675", sa=5)
    table(doc, [["항목", "내용"], ["사용한 AI", ""], ["주고받은 횟수", ""],
                ["AI가 가장 도움이 된 부분", ""], ["AI 결과 중 내가 고친 부분", ""],
                ["사실을 확인한 내용과 방법", ""], ["AI에게 맡기지 않고 내가 한 일", ""]],
          [5.0, 12.0], size=10, rowheight=1.0)

    ev = m.get("guide", {}).get("eval", {})
    if ev.get("rubric"):
        para(doc, "스스로 평가하기", size=15, bold=True, sb=10, sa=2); hr(doc, "DCE3EE")
        para(doc, "선생님이 보는 기준과 같습니다. 제출하기 전에 스스로 표시해 보세요.",
             size=9.5, color="5B6675", sa=4)
        table(doc, [["평가 요소", "상", "중", "하"]] + [[r[0], "☐", "☐", "☐"] for r in ev["rubric"]],
              [10.4, 2.2, 2.2, 2.2], size=10)
    doc.save(path); return path

# ---------------- 교사용 운영 가이드 ----------------
def build_teacher(m, path):
    doc = new_doc(); tc = m.get("guide", {}); ev = tc.get("eval", {})
    head(doc, m["stage"] + " · 공통 과제 모듈 " + m["code"], m["title"] + " — 교사용 운영 가이드",
         m.get("subject", "") + "  |  단계별 AI 수업")

    if tc.get("overview"):
        table(doc, [["구분", "내용"]] + tc["overview"], [3.0, 14.0], size=10)

    h2(doc, "1. 이 모듈의 설계 의도")
    if tc.get("intent"): para(doc, tc["intent"], sa=4)
    if tc.get("intentBox"): box(doc, tc["intentBox"], fill="F7F9FC", border="C9D6EC", size=10)

    h2(doc, "2. 도구 사용 흐름")
    table(doc, [
        ["단계", "무엇을", "누가"],
        ["①", "‘%s’ 파일을 학생 기기에 배포(또는 링크 공유)" % (m["code"] + "_" + m["title"] + ".html"), "교사"],
        ["②", "학생이 이름을 넣고 S1부터 순서대로 진행 (자동 저장됨)", "학생"],
        ["③", "마무리 화면에서 ‘제출 파일 저장(.json)’ 클릭 → 파일 제출", "학생"],
        ["④", "‘과제 채점·피드백 도구(교사용)’에서 .json을 열어 채점·피드백 PDF 회신", "교사"],
    ], [1.6, 12.0, 3.4], size=10)
    para(doc, "※ 인터넷 없이도 동작하며, 브라우저를 닫아도 자동 저장됩니다. "
              "내용을 고치려면 ‘모듈 편집기’에서 수정한 뒤 다시 내보내면 됩니다.", size=9.5, color="5B6675")

    if tc.get("flow"):
        h2(doc, "3. 차시별 운영")
        table(doc, [["차시", "시간", "스텝", "교사가 하는 일"]] + tc["flow"], [1.5, 2.0, 2.2, 11.3], size=9.5)

    doc.add_page_break()
    h2(doc, "4. 스텝별 지도 포인트")
    rows = [["스텝", "핵심", "흔한 실패", "대처"]]
    for st in m["steps"]:
        t = st.get("teach", {})
        rows.append([st["no"] + "\n" + st["title"], t.get("point", ""), t.get("fail", ""), t.get("fix", "")])
    table(doc, rows, [2.2, 4.6, 4.6, 5.6], size=9)

    if tc.get("stuck"):
        doc.add_page_break()
        h2(doc, "5. 학생이 자주 막히는 지점")
        table(doc, [["상황", "대처"]] + tc["stuck"], [4.2, 12.8], size=9.5)

    # ---------- 평가 ----------
    if ev:
        doc.add_page_break()
        h2(doc, "6. 평가 계획")
        if ev.get("outline"): table(doc, [["구분", "내용"]] + ev["outline"], [3.0, 14.0], size=9.5)

        para(doc, "6-1. 평가 요소와 배점", size=12, bold=True, sb=10, sa=3)
        if ev.get("elements"):
            table(doc, [["평가 요소", "관련 스텝", "무엇을 보는가", "배점"]] + ev["elements"],
                  [3.4, 2.4, 9.4, 1.8], size=9)

        para(doc, "6-2. 수준별 채점 기준표", size=12, bold=True, sb=10, sa=3)
        if ev.get("rubric"):
            table(doc, [["평가 요소", "상 (20)", "중 (14)", "하 (8)"]] + ev["rubric"],
                  [3.2, 4.7, 4.6, 4.5], size=8.5)

        doc.add_page_break()
        para(doc, "6-3. 채점표", size=12, bold=True, sa=3)
        para(doc, "① AI와의 협업 · ② 구조와 수정 요청 · ③ 검증 · ④ 나의 기여 · ⑤ 기록과 완성도 (각 20점)",
             size=9, color="5B6675", sa=4)
        cols = ev.get("sheetCols", [])
        n = ev.get("sheetRows", 12)
        if cols:
            widths = [1.2, 2.0, 1.1, 1.1, 1.1, 1.1, 1.1, 1.4, 6.9]
            table(doc, [cols] + [[""] * len(cols) for _ in range(n)], widths, size=9, rowheight=0.85)

        if ev.get("feedback"):
            para(doc, "6-4. 피드백 문구 (골라 쓰기)", size=12, bold=True, sb=10, sa=3)
            table(doc, [["수준", "문구"]] + ev["feedback"], [1.6, 15.4], size=9)

        if ev.get("records"):
            para(doc, "6-5. 관찰 기록 예시 (교과 세부능력 및 특기사항)", size=12, bold=True, sb=10, sa=3)
            para(doc, "그대로 쓰지 말고, 학생의 제출 파일에서 실제로 확인한 내용으로 바꿔 쓰세요.",
                 size=9, color="5B6675", sa=4)
            for i, r in enumerate(ev["records"]):
                box(doc, ["예시 %d" % (i + 1), r], fill="FAFBFD", border="E2E6EC", size=9.5, bold_first=True)

    if tc.get("next"):
        h2(doc, "7. 다음 모듈로 잇기")
        box(doc, tc["next"], fill="F7F9FC", border="C9D6EC", size=10)

    doc.save(path); return path

# ---------------- 학생 배포용 HTML ----------------
def clean(m):
    """저장할 때는 내부용 키(__file)를 뺀다."""
    return {k: v for k, v in m.items() if not k.startswith("__")}

def build_html(m, path, tpl):
    payload = json.dumps(clean(m), ensure_ascii=False).replace("</", "<\\/")
    out = tpl.replace('<script id="moduleData" type="application/json">null</script>',
                      '<script id="moduleData" type="application/json">' + payload + '</script>')
    io.open(path, "w", encoding="utf-8").write(out)
    return path


# ---------------- 모음 파일 · 묶음 JSON ----------------
BUNDLE_HTML = "단계별AI수업_전체모음.html"
BUNDLE_JSON = "_전체모듈.json"

def build_bundle(mods, tpl):
    payload = json.dumps([clean(m) for m in mods], ensure_ascii=False).replace("</", "<\\/")
    out = tpl.replace('<script id="moduleData" type="application/json">null</script>',
                      '<script id="moduleData" type="application/json">' + payload + '</script>')
    p1 = os.path.join(ROOT, BUNDLE_HTML)
    io.open(p1, "w", encoding="utf-8").write(out)
    p2 = os.path.join(MODDIR, BUNDLE_JSON)
    io.open(p2, "w", encoding="utf-8").write(
        json.dumps([clean(m) for m in mods], ensure_ascii=False, indent=1))
    return p1, p2

# ---------------- 전체 관리표 (교사용 워드) ----------------
def ov(m, key, default="—"):
    for row in m.get("guide", {}).get("overview", []):
        if row[0] == key: return row[1]
    return default

def build_admin(mods, path):
    doc = new_doc()
    head(doc, "단계별 AI 수업", "전체 모듈 관리표",
         "3단계 공통 과제 모듈 %d개  |  교사용 한 장 정리" % len(mods))

    para(doc, "이 수업의 전체 흐름", size=13, bold=True, sb=6, sa=3)
    table(doc, [["단계", "내용", "자료"],
                ["1단계", "AI 윤리 — 넣지 말 것 · 믿지 말 것 · 숨기지 말 것", "모듈 0 (10분)"],
                ["2단계", "다양한 AI 맛보기 + RCIF 프롬프트", "미제작"],
                ["3단계", "공통 과제 모듈 — 같이 한 스텝씩", "모듈 %d개 (아래)" % len([x for x in mods if x["code"]!="0"])],
                ["4단계", "각자 주제로 다시 해 보기", "미제작"]],
          [2.0, 10.5, 4.5], size=10)

    para(doc, "모듈 목록", size=13, bold=True, sb=10, sa=3)
    para(doc, "3단계 모듈은 시작 화면에서 ‘짧게 · 1차시(8스텝)’와 ‘전체 · 3차시(14스텝)’ 중 하나를 고를 수 있습니다.",
         size=9.5, color="5B6675", sa=4)
    rows = [["코드", "모듈", "차시", "산출물"]]
    for m in mods:
        rows.append([m["code"], m["title"], ov(m, "차시"), ov(m, "산출물")])
    table(doc, rows, [1.5, 4.2, 2.6, 8.7], size=9)

    para(doc, "모듈마다 다른 ‘장치’", size=13, bold=True, sb=10, sa=3)
    para(doc, "14스텝은 모두 같습니다. 모듈을 가르는 것은 아래 한 가지 조건입니다.",
         size=9.5, color="5B6675", sa=4)
    rows = [["코드", "이 모듈만의 장치"]]
    for m in mods:
        box_ = m.get("guide", {}).get("intentBox", [])
        rows.append([m["code"] + " " + m["title"], box_[0] if box_ else "—"])
    table(doc, rows, [3.6, 13.4], size=9)

    doc.add_page_break()
    para(doc, "모든 모듈이 쓰는 14스텝", size=13, bold=True, sa=3)
    table(doc, [["차시", "스텝", "하는 일"],
                ["1차시", "S1~S2", "재료 정리 → 역할과 상황만 알려주기 (결과물은 아직 요구하지 않음)"],
                ["", "★S3~S4", "AI가 나에게 질문하게 하고, 그 질문에 답하기"],
                ["", "S5~S6", "뼈대만 뽑고 한 번에 하나씩 고쳐 확정"],
                ["2차시", "S7", "1차 초안 — 조건을 촘촘히 걸어서"],
                ["", "S8~S9", "전체 손보기 → 부분만 고치기"],
                ["", "S10", "사실 확인 — 지어낸 곳 찾기"],
                ["3차시", "S11", "상대 입장에서 비판받기"],
                ["", "★S12", "AI 없이 내 손으로 마무리"],
                ["", "S13~S14", "사용 기록표 → 제출하고 짝과 비교"]],
          [1.6, 2.4, 13.0], size=9.5)
    para(doc, "★ S3(AI가 먼저 질문하게)과 S12(AI 없이 내 손으로)가 이 수업의 승부처입니다.",
         size=9.5, color="5B6675")

    para(doc, "권장 운영 순서", size=13, bold=True, sb=10, sa=3)
    box(doc, ["A트랙(문서) — A1 회의록 → A2 선생님께 메일 → A3 학교에 건의 → A4 체험학습 보고서 → A5 근로계약서",
              "B트랙(발표) — B1 발표자료 → B2 대본과 리허설 → B3 예상 질문",
              "낱개로 떼어 써도 됩니다. 처음이라면 A1(회의록)이 가장 쉽습니다.",
              "세 번째 모듈부터는 스텝 설명이 거의 필요 없어집니다."],
        fill="F7F9FC", border="C9D6EC", size=10)

    para(doc, "평가 공통 원칙", size=13, bold=True, sb=10, sa=3)
    box(doc, ["모든 모듈이 5개 요소 × 20점 = 100점 (상 20 · 중 14 · 하 8)입니다.",
              "기준은 ‘AI를 썼는가’가 아니라 ‘기록하고 검증했는가’입니다.",
              "AI 사용은 감점이 아니며, 기록하지 않았거나 확인 없이 제출한 것이 감점입니다.",
              "요소별 세부 기준·채점표·피드백 문구는 모듈별 교사용 운영 가이드에 있습니다."],
        fill="FFF8E6", border="F5DFA0", size=10)

    doc.add_page_break()
    para(doc, "파일과 배포", size=13, bold=True, sa=3)
    table(doc, [["무엇을", "누구에게", "어떻게"],
                ["A○·B○_*.html", "학생", "파일 하나만 주면 됩니다. 인터넷 없이 열리고 자동 저장됩니다."],
                [BUNDLE_HTML, "학생·교사", "모듈 전부가 들어 있는 파일 하나. USB 하나로 여러 모듈 수업이 가능합니다."],
                ["워드/○○_학생용_워크북.docx", "학생", "인쇄해서 나눠 줍니다. 손으로 쓰는 칸이 있습니다."],
                ["워드/○○_교사용_운영가이드.docx", "교사", "차시 운영·지도 포인트·평가 자료"],
                ["모듈편집기.html", "교사", "내용 수정과 새 모듈 만들기, 학생 배포 파일 생성"],
                ["modules/*.json", "교사", "모듈 원본. 편집기로 열어 고칩니다."],
                ["제출 .json", "교사", "‘과제 채점·피드백 도구’에서 열어 채점합니다."]],
          [4.6, 2.4, 10.0], size=9)

    para(doc, "수업 전 점검표", size=13, bold=True, sb=10, sa=3)
    for x in ["학생 기기에서 모듈 파일이 열리는지 한 대로 미리 확인했다",
              "학생들이 쓸 AI(학교 계정)에 접속되는지 확인했다",
              "복사 버튼이 동작하는지 확인했다 (안 되면 드래그 복사 안내)",
              "제출 파일 저장 위치를 학생에게 안내할 준비가 되었다",
              "제출 .json을 채점 도구에서 한 번 열어 봤다",
              "공용 PC라면 이전 학생 기록이 남아 있을 수 있음을 안내한다 (‘처음부터’ 버튼)",
              "워크북을 인쇄했다 (또는 화면으로만 진행하기로 정했다)",
              "이번 시간에 할 모듈과 차시 범위를 정했다"]:
        para(doc, "☐  " + x, size=10, sa=1, indent=0.3)

    doc.save(path); return path

# ---------------- 모듈 편집기 ----------------
def build_editor(tpl):
    src = io.open(os.path.join(ROOT, "_템플릿", "모듈편집기_템플릿.html"), encoding="utf-8").read()
    safe = tpl.replace("</script", r"<\/script")
    out = src.replace("__RUNNER_TEMPLATE__", safe)
    p = os.path.join(ROOT, "모듈편집기.html")
    io.open(p, "w", encoding="utf-8").write(out)
    return p

# ---------------- 허브 ----------------
TRACKS = [("0", "1단계 · 시작 전 약속 (10분)"),
          ("R", "2단계 · AI 고르기와 프롬프트 기본기 (1차시)"),
          ("A", "A트랙 · 문서와 글쓰기"), ("B", "B트랙 · 발표"),
          ("C", "C트랙 · 만들기"), ("D", "D트랙 · 조사와 자료"),
          ("P", "4단계 · 각자 주제로 (마무리)")]

def build_index(mods):
    cards = ""
    basics = sorted([m for m in mods if m.get("basic")], key=lambda x: x["basic"])
    if basics:
        cards += '\n    <h3 class="track">기본 · 먼저 하는 %d가지</h3>' % len(basics)
        cards += _cards(basics)
    for letter, label in TRACKS:
        group = [m for m in mods if m["code"].startswith(letter) and not m.get("basic")]
        if not group:
            continue
        cards += '\n    <h3 class="track">%s</h3>' % label
        cards += _cards(group)
    rest = [m for m in mods if not any(m["code"].startswith(l) for l, _ in TRACKS) and not m.get("basic")]
    if rest:
        cards += '\n    <h3 class="track">그 밖의 모듈</h3>' + _cards(rest)
    return _write_index(cards)

def _cards(mods):
    cards = ""
    for m in mods:
        base = m["code"] + "_" + m["title"].replace(" ", "")
        cards += """
    <div class="mod">
      <div class="mhead"><span class="code">%s</span><h3>%s</h3>%s</div>
      <p class="mdesc">%s</p>
      <div class="links">
        <a class="go" href="%s.html">▶ 학생용 열기</a>
        <a href="워드/%s_학생용_워크북.docx">📄 워크북</a>
        <a href="워드/%s_교사용_운영가이드.docx">📗 교사 가이드</a>
        <a href="modules/%s.json">{ } 모듈 원본</a>
      </div>
    </div>""" % (m["code"], m["title"],
                 ('<span class="basic">기본 %d</span>' % m["basic"]) if m.get("basic") else "",
                 m.get("desc", ""), base, m["code"], m["code"],
                 os.path.splitext(os.path.basename(m["__file"]))[0])
    return cards

def _write_index(cards):
    html = """<!DOCTYPE html>
<html lang="ko"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>단계별 AI 수업</title>
<style>
:root{--ink:#1b2330;--sub:#5b6675;--line:#e2e6ec;--bg:#f4f6fa;--brand:#2f6df0;--brand-d:#1f53c4;--paper:#fff;--soft:#eef3ff;}
*{box-sizing:border-box;}
body{margin:0;font-family:"Apple SD Gothic Neo","Malgun Gothic","맑은 고딕",system-ui,sans-serif;background:var(--bg);color:var(--ink);line-height:1.6;}
.wrap{max-width:860px;margin:0 auto;padding:34px 18px 60px;}
h1{font-size:25px;margin:0 0 6px;letter-spacing:-.5px;}
p.lead{color:var(--sub);margin:0 0 22px;font-size:14px;}
.stages{display:grid;grid-template-columns:repeat(4,1fr);gap:9px;margin-bottom:26px;}
@media(max-width:640px){.stages{grid-template-columns:1fr 1fr;}}
.stage{background:var(--paper);border:1px solid var(--line);border-radius:12px;padding:12px 13px;}
.stage b{display:block;font-size:12px;color:var(--brand-d);margin-bottom:3px;}
.stage span{font-size:12.5px;color:var(--sub);}
.stage.on{border-color:#b9cdf5;background:var(--soft);}
h2{font-size:16px;margin:26px 0 10px;letter-spacing:-.3px;}
.track{font-size:13px;font-weight:800;color:var(--brand-d);background:var(--soft);border:1px solid #c8d8f7;border-radius:8px;padding:6px 11px;display:inline-block;margin:18px 0 10px;}
.mod{background:var(--paper);border:1px solid var(--line);border-radius:14px;padding:16px 18px;margin-bottom:12px;}
.mhead{display:flex;align-items:center;gap:9px;}
.mhead .code{background:linear-gradient(135deg,var(--brand),#7aa2ff);color:#fff;font-weight:900;font-size:12px;padding:3px 9px;border-radius:7px;}
.mhead h3{font-size:17px;margin:0;letter-spacing:-.3px;}
.basic{font-size:11px;font-weight:800;background:#eaf7ec;color:#2c7a44;border:1px solid #bcdcc4;border-radius:20px;padding:2px 9px;margin-left:6px;}
.mdesc{font-size:13.5px;color:var(--sub);margin:7px 0 12px;}
.links{display:flex;flex-wrap:wrap;gap:7px;}
.links a{font-size:12.5px;text-decoration:none;color:var(--sub);border:1px solid var(--line);background:#fbfcfe;border-radius:9px;padding:7px 11px;}
.links a:hover{background:#f2f5fa;}
.links a.go{background:var(--brand);border-color:var(--brand);color:#fff;font-weight:700;}
.tool{display:flex;gap:12px;flex-wrap:wrap;}
.tool a{flex:1 1 220px;background:var(--paper);border:1px solid var(--line);border-radius:12px;padding:14px 16px;text-decoration:none;color:var(--ink);}
.tool a:hover{box-shadow:0 4px 14px rgba(40,60,100,.08);}
.tool b{display:block;font-size:14.5px;margin-bottom:3px;}
.tool span{font-size:12.5px;color:var(--sub);}
.note{font-size:12.5px;color:var(--sub);background:#fff8e6;border:1px solid #f5dfa0;border-radius:11px;padding:11px 13px;margin-top:22px;}
</style></head><body><div class="wrap">
<h1>단계별 AI 수업</h1>
<p class="lead">AI를 처음 쓰는 학생이 한 스텝씩 따라가며 익히는 수업 자료입니다.</p>

<div class="stages">
  <div class="stage"><b>1단계</b><span>AI 윤리 — 넣지 말 것·믿지 말 것·숨기지 말 것</span></div>
  <div class="stage"><b>2단계</b><span>다양한 AI 맛보기 + RCIF 프롬프트</span></div>
  <div class="stage on"><b>3단계</b><span>공통 과제 모듈 — 같이 한 스텝씩</span></div>
  <div class="stage"><b>4단계</b><span>각자 주제로 다시 해 보기</span></div>
</div>

<h2>모듈</h2>%s

<h2>모아 보기 · 관리</h2>
<div class="tool">
  <a href="단계별AI수업_전체모음.html"><b>📚 전체 모음 (파일 하나)</b><span>모듈 전부가 들어 있는 단일 파일. USB 하나로 수업할 수 있고, 진행 상황이 모듈마다 표시됩니다.</span></a>
  <a href="워드/00_전체모듈_관리표.docx"><b>📋 전체 모듈 관리표</b><span>모듈 목록·차시·산출물·장치·평가 원칙·수업 전 점검표 한 장 정리</span></a>
  <a href="modules/_전체모듈.json"><b>{ } 모듈 묶음 파일</b><span>전체 모듈 원본을 한 파일로. 백업과 일괄 편집에 씁니다.</span></a>
  <a href="제미나이_무료버전_수업가이드.html"><b>🟣 무료 제미나이로 수업하기</b><span>학생이 무료 Gemini를 쓸 때 어디가 다른지 · 수업 전 10분 점검 · 모듈별 대안</span></a>
</div>

<h2>교사용 도구</h2>
<div class="tool">
  <a href="모듈편집기.html"><b>🛠 모듈 편집기</b><span>프롬프트·스텝·체크리스트를 화면에서 고치고, 학생 배포용 파일을 만듭니다.</span></a>
  <a href="https://hongyul67-cpu.github.io/ai-task-tools/과제채점_도구_교사용.html" target="_blank" rel="noopener"><b>✅ 과제 채점·피드백 도구</b><span>학생이 낸 .json을 열어 채점하고 피드백 PDF를 만듭니다.</span></a>
</div>

<div class="note"><b>학생에게 나눠 줄 것</b> — 각 모듈의 <b>‘학생용 열기’ 파일 하나</b>면 됩니다.
그 파일 하나에 모든 내용이 들어 있어 인터넷 없이도 열립니다.</div>
</div></body></html>""" % cards
    p = os.path.join(ROOT, "index.html")
    io.open(p, "w", encoding="utf-8").write(html)
    return p

# ---------------- 실행 ----------------
def main():
    tpl = io.open(TPL, encoding="utf-8").read()
    mods = []
    for f in sorted(glob.glob(os.path.join(MODDIR, "*.json"))):
        if os.path.basename(f).startswith("_"):   # 묶음 파일은 모듈이 아님
            continue
        m = json.load(io.open(f, encoding="utf-8")); m["__file"] = f
        mods.append(m)

    for m in mods:
        base = m["code"] + "_" + m["title"].replace(" ", "")
        print(build_html(m, os.path.join(ROOT, base + ".html"), tpl))
        print(build_workbook(m, os.path.join(DOCDIR, m["code"] + "_학생용_워크북.docx")))
        print(build_teacher(m, os.path.join(DOCDIR, m["code"] + "_교사용_운영가이드.docx")))
    b1, b2 = build_bundle(mods, tpl)
    print(b1); print(b2)
    print(build_admin(mods, os.path.join(DOCDIR, "00_전체모듈_관리표.docx")))
    print(build_editor(tpl))
    print(build_index(mods))

    import gemini_guide
    print(gemini_guide.build_docx(os.path.join(DOCDIR, "00_무료제미나이_수업가이드.docx")))
    print(gemini_guide.build_html(os.path.join(ROOT, "제미나이_무료버전_수업가이드.html")))

if __name__ == "__main__":
    main()
